from pathlib import Path

import pymupdf
import pytest
from docx import Document as DocxDocument

from rag_mvp.ingestion.parsers import DocumentParseError, DocumentParser


def test_parses_utf8_text_with_stable_relative_identity(tmp_path: Path) -> None:
    corpus = tmp_path / "corpus"
    source = corpus / "handbooks" / "leave.txt"
    source.parent.mkdir(parents=True)
    source.write_text("Parental leave. 员工育儿假。", encoding="utf-8")
    parser = DocumentParser(corpus_root=corpus)

    first = parser.parse(source)
    second = parser.parse(source)

    assert first.document_id == second.document_id
    assert first.content_sha256 == second.content_sha256
    assert first.source_path == Path("handbooks/leave.txt")
    assert first.pages[0].text == "Parental leave. 员工育儿假。"


def test_rejects_document_outside_corpus(tmp_path: Path) -> None:
    corpus = tmp_path / "corpus"
    corpus.mkdir()
    outside = tmp_path / "outside.txt"
    outside.write_text("outside", encoding="utf-8")

    with pytest.raises(DocumentParseError, match="outside corpus root"):
        DocumentParser(corpus_root=corpus).parse(outside)


def test_parses_native_pdf_pages(tmp_path: Path) -> None:
    source = tmp_path / "architecture.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Internal service architecture and compliance controls.")
    document.save(source)
    document.close()

    parsed = DocumentParser(corpus_root=tmp_path, ocr_enabled=True).parse(source)

    assert len(parsed.pages) == 1
    assert parsed.pages[0].extraction_method == "native"
    assert "service architecture" in parsed.pages[0].text


def test_uses_ocr_fallback_for_text_poor_pdf(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    source = tmp_path / "scan.pdf"
    document = pymupdf.open()
    document.new_page()
    document.save(source)
    document.close()
    parser = DocumentParser(corpus_root=tmp_path, ocr_enabled=True)
    monkeypatch.setattr(parser, "_ocr_pdf_page", lambda _: "扫描的合规指南")

    parsed = parser.parse(source)

    assert parsed.pages[0].extraction_method == "ocr"
    assert parsed.pages[0].text == "扫描的合规指南"


def test_parses_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    source = tmp_path / "specification.docx"
    document = DocxDocument()
    document.add_paragraph("Technical specification")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Timeout"
    table.cell(0, 1).text = "10 seconds"
    document.save(source)

    parsed = DocumentParser(corpus_root=tmp_path).parse(source)

    assert "Technical specification" in parsed.pages[0].text
    assert "Timeout\t10 seconds" in parsed.pages[0].text
