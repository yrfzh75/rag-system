from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path

from rag_mvp.ingestion.models import Document, SourcePage

SUPPORTED_SUFFIXES = frozenset({".pdf", ".docx", ".txt", ".md"})


class DocumentParseError(RuntimeError):
    """Raised when a supported document cannot be parsed."""


class DocumentParser:
    def __init__(
        self,
        *,
        corpus_root: Path,
        ocr_enabled: bool = True,
        ocr_languages: str = "eng+chi_sim",
        ocr_min_text_chars: int = 40,
    ) -> None:
        self.corpus_root = corpus_root.resolve()
        self.ocr_enabled = ocr_enabled
        self.ocr_languages = ocr_languages
        self.ocr_min_text_chars = ocr_min_text_chars

    def parse(self, path: Path) -> Document:
        path = path.resolve()
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            raise DocumentParseError(f"Unsupported file type: {suffix}")

        raw_bytes = path.read_bytes()
        relative_path = self._relative_path(path)
        pages = self._parse_pages(path, suffix)
        if not any(page.text.strip() for page in pages):
            raise DocumentParseError("No text could be extracted")

        return Document(
            document_id=hashlib.sha256(relative_path.encode("utf-8")).hexdigest(),
            content_sha256=hashlib.sha256(raw_bytes).hexdigest(),
            source_path=Path(relative_path),
            source_name=path.name,
            source_type=suffix.removeprefix("."),
            pages=tuple(pages),
        )

    def _relative_path(self, path: Path) -> str:
        try:
            return path.relative_to(self.corpus_root).as_posix()
        except ValueError as exc:
            raise DocumentParseError(f"Document is outside corpus root: {path}") from exc

    def _parse_pages(self, path: Path, suffix: str) -> list[SourcePage]:
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        return [
            SourcePage(
                text=path.read_text(encoding="utf-8-sig"),
                page_number=None,
                extraction_method="native",
            )
        ]

    def _parse_pdf(self, path: Path) -> list[SourcePage]:
        try:
            import pymupdf
        except ImportError as exc:
            raise DocumentParseError("PDF support requires the 'pymupdf' package") from exc

        pages: list[SourcePage] = []
        with pymupdf.open(path) as document:
            for page_index, page in enumerate(document):
                native_text = page.get_text("text", sort=True).strip()
                if len(native_text) >= self.ocr_min_text_chars or not self.ocr_enabled:
                    pages.append(SourcePage(native_text, page_index + 1, "native"))
                    continue
                ocr_text = self._ocr_pdf_page(page).strip()
                pages.append(SourcePage(ocr_text or native_text, page_index + 1, "ocr"))
        return pages

    def _ocr_pdf_page(self, page: object) -> str:
        try:
            import pytesseract
            from PIL import Image
        except ImportError as exc:
            raise DocumentParseError(
                "OCR requires the 'pytesseract' and 'pillow' packages"
            ) from exc

        try:
            pixmap = page.get_pixmap(dpi=200, alpha=False)  # type: ignore[attr-defined]
            image = Image.open(BytesIO(pixmap.tobytes("png")))
            return pytesseract.image_to_string(image, lang=self.ocr_languages)
        except pytesseract.TesseractNotFoundError as exc:
            raise DocumentParseError(
                "OCR was needed but the Tesseract executable is not installed"
            ) from exc
        except pytesseract.TesseractError as exc:
            raise DocumentParseError(
                f"Tesseract OCR failed; verify language packs '{self.ocr_languages}': {exc}"
            ) from exc

    @staticmethod
    def _parse_docx(path: Path) -> list[SourcePage]:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise DocumentParseError("DOCX support requires the 'python-docx' package") from exc

        document = DocxDocument(path)
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append("\t".join(cells))
        return [SourcePage("\n\n".join(blocks), None, "native")]

